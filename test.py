import docx

doc = docx.Document('yufyghj.docx')
# количество абзацев в документе
print(len(doc.paragraphs))

# текст первого абзаца в документе
print(doc.paragraphs[0].text)

# текст второго абзаца в документе
print(doc.paragraphs[1].text)

# текст первого Run второго абзаца
print(doc.paragraphs[1].runs[0].text)